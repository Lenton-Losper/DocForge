"""DOCX file parser."""
from typing import List
from docx import Document as DocxDocument
import re

from models.document_model import Document, Section, Image, DocumentMetadata


def parse_docx(file_path: str, file_name: str) -> Document:
    """
    Parse a DOCX file into normalized Document structure.
    
    Args:
        file_path: Path to the DOCX file
        file_name: Original filename
        
    Returns:
        Normalized Document object
    """
    doc = DocxDocument(file_path)
    sections: List[Section] = []
    images: List[Image] = []
    current_page = 1  # DOCX doesn't have pages, estimate based on content
    
    # Track paragraph content for current section
    current_section_title = None
    current_section_level = 0
    current_section_content: List[str] = []
    
    word_count = 0
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
            
        word_count += len(text.split())
        
        # Check if this is a heading
        is_heading = paragraph.style.name.startswith('Heading')
        
        if is_heading:
            # Save previous section if exists
            if current_section_title is not None:
                sections.append(Section(
                    title=current_section_title,
                    level=current_section_level,
                    content='\n'.join(current_section_content),
                    page=current_page
                ))
            
            # Extract heading level
            heading_match = re.search(r'Heading (\d+)', paragraph.style.name)
            level = int(heading_match.group(1)) if heading_match else 1
            
            # Start new section
            current_section_title = text
            current_section_level = level
            current_section_content = []
        else:
            # Regular paragraph content
            if current_section_title is None:
                # No heading yet, create implicit root section
                current_section_title = "Introduction"
                current_section_level = 1
            current_section_content.append(text)
    
    # Save last section
    if current_section_title is not None:
        sections.append(Section(
            title=current_section_title,
            level=current_section_level,
            content='\n'.join(current_section_content),
            page=current_page
        ))
    
    # Extract images (simplified - docx images are in runs)
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            images.append(Image(
                caption=None,  # Would need to parse nearby text
                page=current_page,
                alt_text=None
            ))
    
    # Estimate page count (rough: ~500 words per page)
    estimated_pages = max(1, word_count // 500)
    
    metadata = DocumentMetadata(
        page_count=estimated_pages,
        word_count=word_count,
        file_type="docx",
        file_name=file_name
    )
    
    return Document(
        sections=sections,
        images=images,
        metadata=metadata
    )
